"""
Word document converter for DocBridge
"""
import re
from typing import Optional, Union, List
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .base import BaseConverter
from ..styles.default import DefaultStyle
from ..formatters.code import CodeFormatter
from ..formatters.table import TableFormatter
from ..formatters.ai_dialog import AIDialogFormatter
from ..utils import parse_ai_dialog, detect_ai_platform, parse_markdown_table


class WordConverter(BaseConverter):
    """
    Convert Markdown to Microsoft Word documents.
    """

    def __init__(self, style: Optional[DefaultStyle] = None):
        super().__init__(style)
        self.code_formatter = CodeFormatter()
        self.table_formatter = TableFormatter()
        self.ai_formatter = AIDialogFormatter()

    def convert(self, markdown_text: str) -> Document:
        """
        Convert markdown text to Word document.

        Args:
            markdown_text: The markdown text to convert

        Returns:
            docx Document object
        """
        # Create new document
        doc = Document()

        # Set up document styles
        self._setup_styles(doc)

        # Detect if this is AI dialog
        ai_platform = detect_ai_platform(markdown_text)

        if ai_platform:
            # Parse as AI dialog
            self._convert_ai_dialog(doc, markdown_text, ai_platform)
        else:
            # Parse as regular markdown
            self._convert_markdown(doc, markdown_text)

        self._document = doc
        return doc

    def convert_file(self, input_path: Union[str, Path],
                     output_path: Optional[Union[str, Path]] = None) -> str:
        """
        Convert markdown file to Word document.

        Args:
            input_path: Path to input markdown file
            output_path: Path for output file (optional)

        Returns:
            Path to output file
        """
        # Read input file
        markdown_text = self._read_file(input_path)

        # Convert to document
        doc = self.convert(markdown_text)

        # Determine output path
        output_path = self._ensure_output_path(input_path, output_path, '.docx')

        # Save document
        doc.save(output_path)

        return str(output_path)

    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles."""
        # Modify Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = self.style.font_name
        font.size = Pt(self.style.font_size_normal)

    def _convert_ai_dialog(self, doc: Document, text: str, platform: str) -> None:
        """Convert AI dialog to document."""
        dialog = self.ai_formatter.parse_dialog(text, platform)

        for entry in dialog:
            # Add speaker paragraph
            p = doc.add_paragraph()
            run = p.add_run(f"{entry['speaker']}: ")
            run.bold = True
            run.font.color.rgb = RGBColor(*entry['color'])

            # Add content
            content_run = p.add_run(entry['content'])

            # Add spacing
            p.paragraph_format.space_after = Pt(self.style.space_paragraph)

    def _convert_markdown(self, doc: Document, text: str) -> None:
        """Convert regular markdown to document."""
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Headings
            if line.startswith('#'):
                self._add_heading(doc, line)
                i += 1
                continue

            # Code blocks
            if line.startswith('```'):
                i = self._add_code_block(doc, lines, i)
                continue

            # Tables
            if '|' in line and self.table_formatter.is_table_line(line):
                i = self._add_table(doc, lines, i)
                continue

            # Regular paragraphs
            self._add_paragraph(doc, line)
            i += 1

    def _add_heading(self, doc: Document, line: str) -> None:
        """Add heading to document."""
        # Count heading level
        level = 0
        for char in line:
            if char == '#':
                level += 1
            else:
                break

        text = line[level:].strip()

        # Add heading
        heading = doc.add_heading(text, level=min(level, 6))

        # Apply style
        style_config = self.style.get_heading_style(min(level, 4))
        for run in heading.runs:
            run.font.size = Pt(style_config['font_size'])
            run.font.bold = style_config.get('bold', False)

    def _add_code_block(self, doc: Document, lines: List[str], start: int) -> int:
        """Add code block to document."""
        # Get language
        first_line = lines[start]
        lang = first_line[3:].strip()

        # Collect code lines
        code_lines = []
        i = start + 1

        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1

        code = '\n'.join(code_lines)

        # Add code block as styled paragraph
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(self.style.space_code_before)
        p.paragraph_format.space_after = Pt(self.style.space_code_after)

        # Add language label if present
        if lang:
            lang_run = p.add_run(f"[{lang}]\n")
            lang_run.italic = True
            lang_run.font.size = Pt(self.style.font_size_small)
            lang_run.font.color.rgb = RGBColor(100, 100, 100)

        # Add code
        code_run = p.add_run(code)
        code_run.font.name = self.style.font_name_mono
        code_run.font.size = Pt(self.style.font_size_code)

        return i + 1

    def _add_table(self, doc: Document, lines: List[str], start: int) -> int:
        """Add table to document."""
        # Collect table lines
        table_lines = []
        i = start

        while i < len(lines) and self.table_formatter.is_table_line(lines[i]):
            table_lines.append(lines[i])
            i += 1

        table_text = '\n'.join(table_lines)
        table_data = self.table_formatter.parse_table(table_text)

        if not table_data['headers']:
            return i

        # Create table
        table = doc.add_table(rows=1, cols=len(table_data['headers']))
        table.style = 'Table Grid'

        # Add headers
        header_cells = table.rows[0].cells
        for j, header in enumerate(table_data['headers']):
            header_cells[j].text = header
            # Style header
            for paragraph in header_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add data rows
        for row_data in table_data['rows']:
            row = table.add_row()
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text

        return i

    def _add_paragraph(self, doc: Document, line: str) -> None:
        """Add regular paragraph to document."""
        p = doc.add_paragraph()

        # Process inline formatting
        self._process_inline_formatting(p, line)

        p.paragraph_format.space_after = Pt(self.style.space_paragraph)

    def _process_inline_formatting(self, paragraph, text: str) -> None:
        """Process inline markdown formatting."""
        # Simple inline formatting
        # Bold: **text** or __text__
        # Italic: *text* or _text_
        # Code: `text`
        # Link: [text](url)

        # For now, add as plain text with basic formatting
        run = paragraph.add_run(text)

        # Apply basic inline formatting
        # This is a simplified version - full implementation would parse more carefully
        if '**' in text or '__' in text:
            run.bold = True
        if '*' in text or '_' in text:
            run.italic = True
